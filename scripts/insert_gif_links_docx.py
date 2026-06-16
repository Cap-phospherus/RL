from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


LINKS = [
    (
        "Random策略完整运行GIF",
        "random_policy_demo.gif",
        Path("D:/C-C++_Examples/RL/outputs/videos/random_policy_demo.gif"),
    ),
    (
        "Greedy策略完整运行GIF",
        "greedy_policy_demo.gif",
        Path("D:/C-C++_Examples/RL/outputs/videos/greedy_policy_demo.gif"),
    ),
    (
        "IPPO sample+shield策略完整运行GIF",
        "ippo_policy_demo.gif",
        Path("D:/C-C++_Examples/RL/outputs/videos/ippo_policy_demo.gif"),
    ),
]


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def add_hyperlink(paragraph, text: str, target: str) -> None:
    r_id = paragraph.part.relate_to(target, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def main() -> int:
    if len(sys.argv) not in {2, 3}:
        print("Usage: insert_gif_links_docx.py <docx_path> [out_docx_path]", file=sys.stderr)
        return 2

    docx_path = Path(sys.argv[1])
    out_path = Path(sys.argv[2]) if len(sys.argv) == 3 else docx_path
    if not docx_path.exists():
        print(f"DOCX not found: {docx_path}", file=sys.stderr)
        return 1

    gif_dir = out_path.parent / "gifs"
    gif_dir.mkdir(exist_ok=True)

    for _, filename, source_path in LINKS:
        if not source_path.exists():
            print(f"GIF not found: {source_path}", file=sys.stderr)
            return 1
        shutil.copy2(source_path, gif_dir / filename)

    backup_path = docx_path.with_name(docx_path.stem + ".before_gif_links.docx")
    if out_path == docx_path and not backup_path.exists():
        shutil.copy2(docx_path, backup_path)

    doc = Document(str(docx_path))
    marker_index = None
    for index, paragraph in enumerate(doc.paragraphs):
        if "三个策略完整的GIF运行视频文件位于" in paragraph.text:
            marker_index = index
            break

    if marker_index is None:
        print("Marker paragraph not found.", file=sys.stderr)
        return 1

    target_indexes = list(range(marker_index + 1, marker_index + 1 + len(LINKS)))
    if target_indexes[-1] >= len(doc.paragraphs):
        print("Not enough paragraphs after marker.", file=sys.stderr)
        return 1

    for paragraph_index, (label, filename, _) in zip(target_indexes, LINKS):
        paragraph = doc.paragraphs[paragraph_index]
        clear_paragraph(paragraph)
        paragraph.add_run("（")
        add_hyperlink(paragraph, label, f"gifs/{filename}")
        paragraph.add_run("）")

    doc.save(str(out_path))
    print(f"Updated: {out_path}")
    if out_path == docx_path:
        print(f"Backup: {backup_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
